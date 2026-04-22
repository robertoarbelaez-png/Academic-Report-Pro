from flask import Flask, render_template, request, jsonify, send_file
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from io import BytesIO
import os
import uuid
from datetime import datetime
import requests
import logging
import re

app = Flask(__name__)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

os.makedirs('informes_generados', exist_ok=True)

DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"

logger.info("=" * 60)
logger.info("🚀 ACADEMIC REPORT PRO - VERSIÓN CORREGIDA")
logger.info(f"🔑 API Key configurada: {'SÍ ✅' if DEEPSEEK_API_KEY else 'NO ❌'}")
logger.info("=" * 60)

# ============================================================
# FUNCIÓN PARA LLAMAR A DEEPSEEK
# ============================================================
def llamar_deepseek(prompt):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "Eres un asistente académico profesional. Generas informes estructurados en español."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 8000,
        "temperature": 0.7
    }
    try:
        response = requests.post(DEEPSEEK_URL, headers=headers, json=data, timeout=120)
        if response.status_code == 200:
            resultado = response.json()
            contenido = resultado['choices'][0]['message']['content']
            contenido = contenido.encode('utf-8', 'ignore').decode('utf-8')
            logger.info(f"✅ Contenido recibido: {len(contenido)} caracteres")
            return contenido
        else:
            logger.error(f"Error HTTP {response.status_code}")
            return None
    except Exception as e:
        logger.error(f"Error: {e}")
        return None

# ============================================================
# EXTRACCIÓN DE SECCIONES - VERSIÓN CORREGIDA
# ============================================================
def extraer_seccion(contenido, nombre):
    """Extrae una sección del contenido - VERSIÓN CORREGIDA"""
    
    if not contenido:
        return ""
    
    # Limpiar caracteres basura
    contenido = contenido.encode('utf-8', 'ignore').decode('utf-8')
    contenido = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', contenido)
    
    # Buscar el título en diferentes formatos
    patrones_titulo = [
        rf'\*\*{nombre}\*\*:?',
        rf'\*\*{nombre}\*\*',
        rf'### {nombre}',
        rf'#{nombre}#',
        rf'{nombre}',
    ]
    
    for patron_titulo in patrones_titulo:
        patron_seccion = rf'{patron_titulo}\s*(.*?)(?=\n\s*\*\*[A-ZÁÉÍÓÚÜÑ]|\n\s*\d+\.\s*\*\*|\Z)'
        match = re.search(patron_seccion, contenido, re.DOTALL | re.IGNORECASE)
        if match:
            texto = match.group(1).strip()
            if len(texto) > 100:
                texto = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', texto)
                texto = texto.replace('\n', '<br/>')
                return texto
    
    # Búsqueda más simple
    idx = contenido.upper().find(nombre.upper())
    if idx != -1:
        siguiente_match = re.search(r'\n\s*\*\*[A-ZÁÉÍÓÚÜÑ]|\n\s*\d+\.\s*\*\*', contenido[idx+len(nombre):])
        if siguiente_match:
            texto = contenido[idx+len(nombre):idx+len(nombre)+siguiente_match.start()].strip()
        else:
            texto = contenido[idx+len(nombre):].strip()
        
        if len(texto) > 100:
            texto = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', texto)
            texto = texto.replace('\n', '<br/>')
            return texto
    
    return ""

# ============================================================
# PROMPTS DEDICADOS POR SECCIÓN
# ============================================================
PROMPTS_SECCIONES = {
    'introduccion': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE la sección de INTRODUCCIÓN para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}. Información adicional: {info or 'Ninguna'}.

Requisitos:
- Mínimo 4 párrafos completos y bien desarrollados
- Incluir: contexto general del tema, justificación e importancia, planteamiento del problema, estructura del informe
- Redacción formal y académica
- NO incluyas el título, solo el contenido

Escribe el contenido directamente, sin títulos ni encabezados.""",

    'objetivos': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE la sección de OBJETIVOS para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}.

Formato exacto:
OBJETIVO GENERAL:
[Un objetivo general claro y medible que abarque todo el informe]

OBJETIVOS ESPECÍFICOS:
1. [Objetivo específico 1 — acción concreta y verificable]
2. [Objetivo específico 2 — acción concreta y verificable]
3. [Objetivo específico 3 — acción concreta y verificable]
4. [Objetivo específico 4 — acción concreta y verificable]
5. [Objetivo específico 5 — acción concreta y verificable]

Usa verbos en infinitivo (Analizar, Identificar, Evaluar, Determinar, Describir...).
Escribe el contenido directamente, sin título de sección.""",

    'marco_teorico': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE el MARCO TEÓRICO para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}. Información adicional: {info or 'Ninguna'}.

Requisitos:
- Mínimo 5 párrafos completos
- Incluir: antecedentes históricos del tema, definiciones de conceptos clave, teorías y modelos relevantes, estado actual del conocimiento
- Citar autores y fuentes relevantes de manera académica
- Redacción formal con vocabulario especializado

Escribe el contenido directamente, sin título de sección.""",

    'metodologia': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE la sección de METODOLOGÍA para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}.

Requisitos:
- Mínimo 4 párrafos completos
- Incluir: tipo y enfoque de investigación, población o muestra (si aplica), técnicas e instrumentos de recolección de datos, procedimiento de análisis, consideraciones éticas
- Justificar cada decisión metodológica
- Redacción en pasado o presente académico

Escribe el contenido directamente, sin título de sección.""",

    'desarrollo': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE el DESARROLLO para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}. Información adicional: {info or 'Ninguna'}.

Requisitos:
- Mínimo 6 párrafos completos — esta es la sección más extensa
- Incluir: presentación de resultados o hallazgos, análisis detallado, discusión crítica, comparación con teorías del marco teórico, datos relevantes o ejemplos concretos
- Organizar con subtemas claros dentro del contenido
- Profundidad académica real

Escribe el contenido directamente, sin título de sección.""",

    'conclusiones': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE las CONCLUSIONES para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}.

Formato:
1. [Conclusión 1: responde directamente al objetivo general — 3 oraciones mínimo]
2. [Conclusión 2: hallazgo más importante del desarrollo — 3 oraciones mínimo]
3. [Conclusión 3: implicaciones prácticas o teóricas — 3 oraciones mínimo]
4. [Conclusión 4: limitaciones encontradas — 2 oraciones mínimo]
5. [Conclusión 5: perspectivas futuras o reflexión final — 3 oraciones mínimo]

Cada conclusión debe ser un párrafo sustancial, no solo una oración.
Escribe el contenido directamente, sin título de sección.""",

    'recomendaciones': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE las RECOMENDACIONES para un informe académico de tipo "{tipo}" sobre: "{tema}".
Nivel educativo: {nivel}.

Formato:
1. [Recomendación 1: dirigida a quién y qué acción concreta — 3 oraciones]
2. [Recomendación 2: dirigida a quién y qué acción concreta — 3 oraciones]
3. [Recomendación 3: dirigida a quién y qué acción concreta — 3 oraciones]
4. [Recomendación 4: dirigida a quién y qué acción concreta — 3 oraciones]
5. [Recomendación 5: para futuras investigaciones — 3 oraciones]

Cada recomendación debe ser práctica, específica y justificada.
Escribe el contenido directamente, sin título de sección.""",

    'referencias': lambda tema, info, tipo, norma, nivel: f"""Escribe ÚNICAMENTE las REFERENCIAS para un informe académico sobre: "{tema}".
Formato de citación: {norma}. Nivel: {nivel}.

Genera 8 referencias bibliográficas reales y relevantes para este tema en formato {norma} estricto.
- Incluir libros, artículos de revista, reportes de organismos y fuentes web académicas
- Ordenadas alfabéticamente por apellido del autor
- Años entre 2015 y 2024 preferiblemente
- Autores, títulos y editoriales plausibles y académicos

Escribe solo la lista de referencias, sin título de sección."""
}

# ============================================================
# GENERAR UNA SECCIÓN INDIVIDUAL
# ============================================================
def generar_seccion(seccion, tema, info_extra, tipo_informe, norma, nivel):
    """Genera una sola sección con un prompt dedicado"""
    if seccion not in PROMPTS_SECCIONES:
        return None

    prompt = PROMPTS_SECCIONES[seccion](tema, info_extra, tipo_informe, norma, nivel)
    system_prompt = (
        "Eres un experto en redacción académica universitaria en español. "
        "Escribes contenido sustancial, formal y bien estructurado. "
        "Respondes SOLO con el contenido solicitado, sin títulos de sección, sin preámbulos, sin comentarios adicionales."
    )

    contenido = llamar_deepseek(prompt, system_prompt=system_prompt, max_tokens=3000)

    if not contenido:
        return None

    contenido = contenido.strip()
    contenido = contenido.replace('\n', '<br/>')
    logger.info(f"Sección '{seccion}' generada: {len(contenido)} caracteres")
    return contenido


# ============================================================
# GENERAR INFORME COMPLETO (mantener para compatibilidad)
# ============================================================
def generar_informe_completo(tema, info_extra, tipo_informe, norma, nivel):
    """Genera todas las secciones en secuencia — usado como fallback"""
    secciones = {}
    claves = ['introduccion', 'objetivos', 'marco_teorico', 'metodologia',
              'desarrollo', 'conclusiones', 'recomendaciones', 'referencias']
    for clave in claves:
        resultado = generar_seccion(clave, tema, info_extra, tipo_informe, norma, nivel)
        secciones[clave] = resultado or ''
    return secciones

# ============================================================
# GENERAR PDF
# ============================================================
def generar_pdf(datos_usuario, secciones):
    nombre = datos_usuario.get('nombre', 'Estudiante')
    tema = datos_usuario.get('tema', 'Tema')
    asignatura = datos_usuario.get('asignatura', '')
    profesor = datos_usuario.get('profesor', '')
    institucion = datos_usuario.get('institucion', '')
    ciudad = datos_usuario.get('ciudad', '')
    fecha = datos_usuario.get('fecha', datetime.now().strftime('%d/%m/%Y'))
    norma = datos_usuario.get('norma', 'APA 7')
    
    filename = f"informe_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.pdf"
    filepath = os.path.join('informes_generados', filename)
    
    styles = getSampleStyleSheet()
    
    styles.add(ParagraphStyle(name='TextoJustificado', parent=styles['Normal'], alignment=TA_JUSTIFY, fontSize=11, fontName='Times-Roman', spaceAfter=12, leading=16))
    styles.add(ParagraphStyle(name='Titulo1', parent=styles['Heading1'], fontSize=16, fontName='Helvetica-Bold', textColor=colors.HexColor('#1a365d'), spaceBefore=24, spaceAfter=12))
    styles.add(ParagraphStyle(name='TituloPortada', parent=styles['Title'], fontSize=24, alignment=TA_CENTER, textColor=colors.HexColor('#1a365d')))
    styles.add(ParagraphStyle(name='TextoCentrado', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12))
    
    doc = SimpleDocTemplate(filepath, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    story = []
    
    # Portada
    story.append(Spacer(1, 2.0*inch))
    story.append(Paragraph("INFORME ACADÉMICO", styles['TituloPortada']))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(tema.upper(), styles['TextoCentrado']))
    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph(f"<b>Presentado por:</b> {nombre}", styles['TextoCentrado']))
    if asignatura:
        story.append(Paragraph(f"<b>Asignatura:</b> {asignatura}", styles['TextoCentrado']))
    if profesor:
        story.append(Paragraph(f"<b>Docente:</b> {profesor}", styles['TextoCentrado']))
    if institucion:
        story.append(Paragraph(f"<b>Institución:</b> {institucion}", styles['TextoCentrado']))
    if ciudad:
        story.append(Paragraph(f"<b>Ciudad:</b> {ciudad}", styles['TextoCentrado']))
    story.append(Paragraph(f"<b>Fecha:</b> {fecha}", styles['TextoCentrado']))
    story.append(Paragraph(f"<b>Norma:</b> {norma}", styles['TextoCentrado']))
    story.append(PageBreak())
    
    # Índice
    story.append(Paragraph("ÍNDICE", styles['Titulo1']))
    story.append(Spacer(1, 0.2*inch))
    for idx in ["1. INTRODUCCIÓN", "2. OBJETIVOS", "3. MARCO TEÓRICO", "4. METODOLOGÍA", "5. DESARROLLO", "6. CONCLUSIONES", "7. RECOMENDACIONES", "8. REFERENCIAS"]:
        story.append(Paragraph(f"• {idx}", styles['TextoJustificado']))
        story.append(Spacer(1, 0.1*inch))
    story.append(PageBreak())
    
    # Secciones
    secciones_orden = [
        ("1. INTRODUCCIÓN", 'introduccion'),
        ("2. OBJETIVOS", 'objetivos'),
        ("3. MARCO TEÓRICO", 'marco_teorico'),
        ("4. METODOLOGÍA", 'metodologia'),
        ("5. DESARROLLO", 'desarrollo'),
        ("6. CONCLUSIONES", 'conclusiones'),
        ("7. RECOMENDACIONES", 'recomendaciones'),
        ("8. REFERENCIAS", 'referencias')
    ]
    
    for titulo, clave in secciones_orden:
        story.append(Paragraph(titulo, styles['Titulo1']))
        story.append(Spacer(1, 0.2*inch))
        contenido = secciones.get(clave, '')
        if contenido and len(contenido) > 100:
            story.append(Paragraph(contenido, styles['TextoJustificado']))
        else:
            story.append(Paragraph("No se pudo generar esta sección.", styles['TextoJustificado']))
        story.append(PageBreak())
    
    doc.build(story)
    return filename, filepath

# ============================================================
# GENERAR WORD
# ============================================================
def generar_word(datos_usuario, secciones):
    nombre = datos_usuario.get('nombre', 'Estudiante')
    tema = datos_usuario.get('tema', 'Tema')
    asignatura = datos_usuario.get('asignatura', '')
    profesor = datos_usuario.get('profesor', '')
    institucion = datos_usuario.get('institucion', '')
    ciudad = datos_usuario.get('ciudad', '')
    fecha = datos_usuario.get('fecha', datetime.now().strftime('%d/%m/%Y'))
    norma = datos_usuario.get('norma', 'APA 7')
    
    doc = Document()
    doc.add_heading('INFORME ACADÉMICO', 0)
    doc.add_heading(tema, level=1)
    doc.add_paragraph(f"Presentado por: {nombre}")
    if asignatura: doc.add_paragraph(f"Asignatura: {asignatura}")
    if profesor: doc.add_paragraph(f"Docente: {profesor}")
    if institucion: doc.add_paragraph(f"Institución: {institucion}")
    if ciudad: doc.add_paragraph(f"Ciudad: {ciudad}")
    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph(f"Norma: {norma}")
    doc.add_page_break()
    
    doc.add_heading('ÍNDICE', level=2)
    for titulo in ['INTRODUCCIÓN', 'OBJETIVOS', 'MARCO TEÓRICO', 'METODOLOGÍA', 'DESARROLLO', 'CONCLUSIONES', 'RECOMENDACIONES', 'REFERENCIAS']:
        doc.add_paragraph(f"• {titulo}")
    doc.add_page_break()
    
    secciones_orden = [
        ("INTRODUCCIÓN", 'introduccion'), ("OBJETIVOS", 'objetivos'),
        ("MARCO TEÓRICO", 'marco_teorico'), ("METODOLOGÍA", 'metodologia'),
        ("DESARROLLO", 'desarrollo'), ("CONCLUSIONES", 'conclusiones'),
        ("RECOMENDACIONES", 'recomendaciones'), ("REFERENCIAS", 'referencias')
    ]
    
    for titulo, clave in secciones_orden:
        doc.add_heading(titulo, level=2)
        contenido = secciones.get(clave, '')
        if contenido:
            doc.add_paragraph(contenido.replace('<br/>', '\n').replace('<b>', '').replace('</b>', ''))
        else:
            doc.add_paragraph("No se pudo generar esta sección.")
        doc.add_page_break()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================
# RUTAS DE LA API
# ============================================================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['POST'])
def generar():
    try:
        data = request.json
        tema = data.get('tema', '').strip()
        nivel = data.get('nivel', 'universitario')
        modo = data.get('modo', 'rapido')
        tipo_informe = data.get('tipo_informe', 'academico')
        norma = data.get('norma', 'APA 7')
        nombre = data.get('nombre', 'Estudiante')
        asignatura = data.get('asignatura', '')
        profesor = data.get('profesor', '')
        institucion = data.get('institucion', '')
        ciudad = data.get('ciudad', '')
        texto_usuario = data.get('texto_usuario', '')
        
        autores = data.get('autores', [])
        if autores:
            nombre_principal = autores[0].get('nombre', nombre)
        else:
            nombre_principal = nombre
        
        if not tema:
            return jsonify({'success': False, 'error': 'El tema es requerido'}), 400
        
        logger.info(f"📨 Generando informe - Tema: {tema[:50]}...")
        
        secciones = generar_informe_completo(tema, texto_usuario, tipo_informe, norma, nivel)
        
        if not secciones:
            return jsonify({'success': False, 'error': 'No se pudo generar el informe'}), 500
        
        datos_usuario = {
            'nombre': nombre_principal,
            'tema': tema,
            'asignatura': asignatura,
            'profesor': profesor,
            'institucion': institucion,
            'ciudad': ciudad,
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'norma': norma
        }
        
        return jsonify({'success': True, 'secciones': secciones, 'datos_usuario': datos_usuario})
    except Exception as e:
        logger.error(f"Error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/generar-seccion', methods=['POST'])
def generar_seccion_endpoint():
    """Genera una sola sección — llamado 8 veces desde el frontend"""
    try:
        data = request.json
        seccion   = data.get('seccion', '')
        tema      = data.get('tema', '').strip()
        nivel     = data.get('nivel', 'universitario')
        tipo      = data.get('tipo_informe', 'academico')
        norma     = data.get('norma', 'APA 7')
        info      = data.get('texto_usuario', '')

        if not seccion or not tema:
            return jsonify({'success': False, 'error': 'Faltan parámetros'}), 400

        logger.info(f"Generando sección '{seccion}' para tema: {tema[:40]}...")
        contenido = generar_seccion(seccion, tema, info, tipo, norma, nivel)

        if contenido:
            return jsonify({'success': True, 'seccion': seccion, 'contenido': contenido})
        else:
            return jsonify({'success': False, 'error': f'No se pudo generar: {seccion}'}), 500

    except Exception as e:
        logger.error(f"Error en /generar-seccion: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/exportar-pdf', methods=['POST'])
def exportar_pdf():
    data = request.json
    filename, filepath = generar_pdf(data['datos_usuario'], data['secciones'])
    return send_file(filepath, as_attachment=True, download_name=filename)

@app.route('/exportar-word', methods=['POST'])
def exportar_word():
    data = request.json
    buffer = generar_word(data['datos_usuario'], data['secciones'])
    return send_file(buffer, as_attachment=True, download_name=f"informe_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

@app.route('/preview', methods=['POST'])
def preview():
    try:
        data = request.json
        tema = data.get('tema', '')
        prompt = f"Genera un breve resumen sobre: {tema} en 300 palabras"
        contenido = llamar_deepseek(prompt)
        if contenido:
            return jsonify({'success': True, 'contenido': contenido[:1000]})
        return jsonify({'success': False, 'error': 'No se pudo generar'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/health')
def health():
    return jsonify({'status': 'healthy', 'api_configured': bool(DEEPSEEK_API_KEY)})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)
